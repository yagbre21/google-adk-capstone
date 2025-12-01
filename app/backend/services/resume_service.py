"""
Resume extraction and processing service.
"""
import io
from typing import Optional
import PyPDF2
import docx


class ResumeExtractionError(Exception):
    """Raised when resume extraction fails."""
    pass


class ResumeService:
    """Service for extracting and processing resume content."""

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md'}
    MIN_RESUME_LENGTH = 100

    def extract_from_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF file."""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text_parts = []

            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            text = "\n".join(text_parts).strip()

            if not text:
                raise ResumeExtractionError("Could not extract text from PDF. The file may be image-based.")

            return text

        except PyPDF2.errors.PdfReadError as e:
            raise ResumeExtractionError(f"Invalid PDF file: {str(e)}")

    def extract_from_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            text = "\n".join(text_parts).strip()

            if not text:
                raise ResumeExtractionError("Could not extract text from DOCX. The file may be empty.")

            return text

        except Exception as e:
            raise ResumeExtractionError(f"Invalid DOCX file: {str(e)}")

    def extract_from_file(self, file_bytes: bytes, filename: str) -> str:
        """Extract text from file based on extension."""
        filename_lower = filename.lower()

        if filename_lower.endswith('.pdf'):
            return self.extract_from_pdf(file_bytes)
        elif filename_lower.endswith('.docx'):
            return self.extract_from_docx(file_bytes)
        elif filename_lower.endswith('.txt') or filename_lower.endswith('.md'):
            return self.extract_from_text(file_bytes)
        else:
            raise ResumeExtractionError(
                f"Unsupported file type. Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

    def extract_from_text(self, file_bytes: bytes) -> str:
        """Extract text from TXT or MD file."""
        try:
            text = file_bytes.decode('utf-8').strip()
            if not text:
                raise ResumeExtractionError("Text file is empty.")
            return text
        except UnicodeDecodeError:
            raise ResumeExtractionError("Could not decode text file. Please ensure it's UTF-8 encoded.")

    def validate_text(self, text: str) -> str:
        """Validate and clean resume text."""
        if not text:
            raise ResumeExtractionError("Resume text is empty.")

        cleaned = text.strip()

        if len(cleaned) < self.MIN_RESUME_LENGTH:
            raise ResumeExtractionError(
                f"Resume text is too short ({len(cleaned)} chars). "
                f"Please provide a complete resume (min {self.MIN_RESUME_LENGTH} chars)."
            )

        return cleaned
